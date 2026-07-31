# -*- coding: utf-8 -*-
"""生成《参加挑战赛理由及作品亮点》Word 文档。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# ---- 全局字体：正文宋体 / 标题黑体，西文用兼容字体 ----
def set_run_font(run, name_cn="宋体", name_en="Times New Roman", size=12, bold=False, color=None):
    run.font.name = name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name_cn)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

# 设置默认样式
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(12)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, name_cn="黑体", size=20, bold=True, color=(0x1F, 0x3B, 0x73))
    p.space_after = Pt(6)

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, name_cn="楷体", size=12, color=(0x55, 0x55, 0x55))

def add_h1(text):
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    p.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, name_cn="黑体", size=15, bold=True, color=(0x1F, 0x3B, 0x73))

def add_h2(text):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, name_cn="黑体", size=13, bold=True, color=(0x2E, 0x5C, 0x8A))

def add_body(text, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=12)
    return p

def add_bullet(text, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r = p.add_run(bold_lead)
        set_run_font(r, size=12, bold=True, color=(0x1F, 0x3B, 0x73))
        r2 = p.add_run(text)
        set_run_font(r2, size=12)
    else:
        r = p.add_run(text)
        set_run_font(r, size=12)

# ============ 封面 ============
add_title("参加挑战赛理由及作品亮点")
add_subtitle("——灵山小向导·灵曦  景区导览服务 AI 数字人系统")
add_subtitle("第十五届中国软件杯  A 组赛题 A5")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("2026 年 8 月")
set_run_font(r, name_cn="楷体", size=12, color=(0x55, 0x55, 0x55))

# ============ 一、参赛理由 ============
add_h1("一、参加挑战赛的理由")

add_h2("1. 契合国家文旅数字化转型战略，解决真实行业痛点")
add_body(
    "国家大力推进文旅产业数字化转型，智慧景区建设已成为行业趋势。然而当前景区导览服务普遍存在"
    "导游资源稀缺、信息单向传递、缺乏情感连接和管理反馈盲区四大痛点：旺季专业导游供不应求，"
    "传统录音导览设备内容固定且无法互动，冰冷设备难以提供亲切感，景区管理者更难以量化评估服务"
    "质量。本赛题来源于锐捷网络的真实企业业务场景，我们选择参赛，正是希望用 AI 数字人技术直面"
    "这些痛点，构建一个 7×24 小时在线、可交互、有温度、可量化的“智能导游”，把课堂上的大模型、"
    "多模态与检索增强生成技术真正落到一个能用的产品里。"
)

add_h2("2. 以赛促学，打通“大模型—多模态—数字人”全栈工程能力")
add_body(
    "本赛题涉及数字人建模与驱动、语音识别与合成、自然语言处理、大模型技术和多模态情感分析等"
    "多个前沿方向，是一次难得的全栈工程锻炼。我们希望借参赛机会，从零搭建一条从语音输入到数字人"
    "口型同步播报的完整链路，深入理解 RAG 检索增强、流式生成、端到端延迟优化、GPU 资源调度与"
    "多级降级等工程关键点，把分散的知识整合成一套可部署、可评测、可演进的真实系统，而非停留在"
    "调用 API 的演示层面。"
)

add_h2("3. 践行“国产 + 开源”技术路线，掌握自主可控能力")
add_body(
    "赛题明确建议使用开源或国产技术。我们在方案中以智谱 GLM 系列多模态大模型为核心，配合开源的"
    "BGE-M3 向量模型、FAISS 检索、Whisper/ParaTTS 思路的语音方案、LiveTalking 与 Wav2Lip 数字人"
    "驱动、HumanOmni 情绪分析以及 Qwen2 本地模型，形成“云端 GLM + 本地开源模型”双路线可切换"
    "架构。参赛过程让我们切身体会国产模型的能力边界与开源生态的工程化方法，为后续自主可控的"
    "产业落地积累经验。"
)

add_h2("4. 建立数据闭环，体现工程严谨与可验证性")
add_body(
    "我们重视“可验证”而非“看起来能用”。系统建立了冻结事实评测集、延迟实测方法、服务健康检查和"
    "资源回收机制，并把准确率与延迟等指标与赛题要求逐项对应。参赛促使我们以工程师标准对待每一"
    "个数字，养成“指标可追溯、结论有依据”的职业习惯。"
)

# ============ 二、作品亮点 ============
add_h1("二、作品亮点")

add_h2("1. 可靠的景区知识问答：本地 RAG + 引用追踪 + 拒答约束")
add_bullet(
    "以本地景区资料为唯一事实依据，BGE-M3 生成向量、FAISS 检索，回答附带知识引用片段，"
    "资料不足时主动拒答，杜绝“无依据的编造”。",
    bold_lead="事实可追溯："
)
add_bullet(
    "已建立 85 题冻结评测集，覆盖事实/路线、同义问法、模糊问法、跨景区混淆、无资料拒答、"
    "知识冲突与提示注入等场景；15 题事实冒烟基线为 15/15，并明确说明其非最终准确率，"
    "赛前将用不少于 80 道专家复核题产出分类评测报告。",
    bold_lead="可评测："
)
add_bullet(
    "问答可在云端 GLM API 与本地 Qwen2-7B 之间切换，弱网或离线环境下仍可保障基础问答。",
    bold_lead="双模型可切换："
)

add_h2("2. 自然的语音数字人：语义断句 + 流式播报 + 口型同步")
add_bullet(
    "模型回答采用 SSE 流式返回，FastAPI 按完整语义句切分并并行衔接 TTS 与 LiveTalking，"
    "不必等全文生成完才开始播报，实现“首句即说”。",
    bold_lead="首句即说："
)
add_bullet(
    "通过 LiveTalking + Wav2Lip FP16 驱动数字人口型与表情同步；新问题到达时取消旧播报队列，"
    "防止音频重叠，保证讲解自然连贯。",
    bold_lead="口型与表情同步："
)
add_bullet(
    "实测“发送文字 → 首个非静音音频”中位数约 2.99 秒，“语音输入结束 → 首个非静音数字人音频帧”"
    "约 3.37 秒，均低于赛题 5 秒要求，并保留完整复测方法与分阶段数据。",
    bold_lead="延迟达标："
)

add_h2("3. 多模态交互与个性化：语音/文字/图片/情绪 + 智能路线")
add_bullet(
    "支持语音、文字、拍照识景多种输入；拍照识景采用 CLIP 景点图集召回 + 参考图复核反证，"
    "再由 RAG 对确认景点提供有依据讲解，降低视觉误判风险。",
    bold_lead="多模态输入："
)
add_bullet(
    "根据游客兴趣（历史/自然/亲子）、时长（2/4/6 小时）、同行人群与步行偏好生成个性化路线，"
    "在地图上绘制编号点位与游览顺序，并提供逐站行程和高德分段步行导航入口。",
    bold_lead="个性化路线："
)
add_bullet(
    "游客侧情绪分析采用音频—文本多模态（不采集人脸视频，兼顾隐私），数字人表情按情绪策略实时切换。",
    bold_lead="情感互动："
)

add_h2("4. 弱定位场景的工程化解决：多级降级与地图兜底")
add_body(
    "针对赛题可选的“GPS 信号弱或难以定位”现实场景，系统设计了多级定位降级：GPS → 景点码 → "
    "手动选择；景区无可核验点位时返回 422 并提示；高德底图不可用时仍保留文字路线、逐站行程与"
    "导航入口。路线点位保留坐标系、来源、测绘状态与预计精度，游览顺序线明确标注“非道路级导航”，"
    "杜绝数据真实性误导。"
)

add_h2("5. 完整的管理运营闭环：知识库 + 数字人配置 + 数据大屏")
add_bullet(
    "管理员可上传、更新、删除知识文档并触发索引重建，索引重建成功后才替换当前索引，保障稳定性。",
    bold_lead="知识库管理："
)
add_bullet(
    "可配置数字人显示名称、已安装形象与合成声音，表情由情绪策略实时驱动。",
    bold_lead="数字人形象管理："
)
add_bullet(
    "运营数据大屏展示服务人次、热门问答、响应时间、情感趋势、满意度与服务建议，"
    "无评分或无情绪结果时显示“暂无”，不生成虚假指标。",
    bold_lead="数据大屏与感受度报告："
)

add_h2("6. 工程质量与资源治理：分层架构 + GPU 按需调度 + 安全合规")
add_bullet(
    "FastAPI 网关统一编排，RAG、GLM 多模态、LiveTalking、HumanOmni 作为独立服务，"
    "模块独立部署、独立健康检查，单点故障可降级而不阻塞基础问答。",
    bold_lead="分层可维护："
)
add_bullet(
    "LiveTalking 权重常驻 CPU/RAM，建立 WebRTC 会话后才选空闲 GPU 上显存（实测约占 756MiB），"
    "页面关闭或心跳超时后自动移回 CPU；本地 7B 模型空闲 120 秒自动卸载，按需释放显存。",
    bold_lead="GPU 按需调度："
)
add_bullet(
    "管理接口鉴权 + 来源校验，模型密钥仅服务端保存不下发浏览器，上传文件校验类型/大小/路径，"
    "多模态原始媒体分析后删除只留结构化结果，麦克风访问使用 HTTPS 安全上下文。",
    bold_lead="安全合规："
)

add_h2("7. 多端可达与真实可部署")
add_body(
    "系统提供游客网页端（文字 / 麦克风）、景区管理后台、OpenAPI 文档与安卓 APK，"
    "采用“本地 GPU 计算节点 + 云端公网转发节点”的演示部署拓扑，可现场打开真实游客端与管理端"
    "演示，而非仅放录屏。该拓扑可迁移至景区私有服务器或云端 GPU 环境，具备真实落地条件。"
)

# ============ 三、总结 ============
add_h1("三、总结")
add_body(
    "“灵山小向导·灵曦”紧扣赛题痛点，以国产多模态大模型为核心、本地 RAG 知识库为事实根基、"
    "开源数字人方案为交互载体，构建了一个多模态可交互、知识可信赖、延迟达标、弱定位可降级、"
    "管理可量化的景区导览 AI 数字人系统。作品既体现了大模型与多模态技术的集成创新，也保持了"
    "工程严谨与可验证性，是一套面向真实文旅场景、具备落地与推广价值的解决方案。"
)

out = "/home/gmn/codes/cup/deliverables/A5-灵山小向导-参赛理由及作品亮点.docx"
doc.save(out)
print("saved:", out)
