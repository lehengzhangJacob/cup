# rag/chunker.py
import json
import re
import pandas as pd
from docx import Document
from rag.config import (
    GUIDELINE_DOCX, DATASET_DOCX, EXTRA_DOCS_DIR, XLSX_FILE, WHITELIST_JSON,
    MAX_CHUNK_CHARS, CHUNK_OVERLAP
)


def _load_whitelist_full():
    return json.loads(WHITELIST_JSON.read_text(encoding="utf-8"))


def _split_sentences(text):
    """按中文句末标点/换行切句，保留标点。用于在句子边界处分块，避免断句。"""
    parts = re.split(r"(?<=[。！？\n])", text)
    return [s for s in (p.strip() for p in parts) if s]


def _split_long_text(text, max_chars=MAX_CHUNK_CHARS, overlap=CHUNK_OVERLAP):
    """把超长文本按句子边界切成多块，块间保留整句重叠，避免从词/句中间切断。"""
    if not text or not text.strip():
        return []
    if len(text) <= max_chars:
        return [text.strip()]

    sentences = _split_sentences(text)
    chunks, current = [], []
    cur_len = 0
    for sent in sentences:
        # 单句本身超长：按字符硬切（极少见）
        if len(sent) > max_chars:
            if current:
                chunks.append("".join(current).strip())
                current, cur_len = [], 0
            for i in range(0, len(sent), max_chars - overlap):
                chunks.append(sent[i:i + max_chars])
            continue
        if cur_len + len(sent) > max_chars and current:
            chunks.append("".join(current).strip())
            # 用末尾整句做重叠（不超过overlap长度）
            tail, tail_len = [], 0
            for s in reversed(current):
                if tail_len + len(s) > overlap and tail:
                    break
                tail.insert(0, s)
                tail_len += len(s)
            current = tail
            cur_len = tail_len
        current.append(sent)
        cur_len += len(sent)
    if current:
        chunks.append("".join(current).strip())
    return [c for c in chunks if c] or [text[:max_chars].strip()]


def _match_attraction(value, names, aliases):
    """Match an xlsx attraction_name value to a canonical 灵山胜境 whitelist name.

    严格匹配：只保留确属灵山胜境的景点，避免把无关景点（如无锡梁溪区的“南禅寺”）
    通过“禅寺”这类单字别名误并入“祥符禅寺”。

    匹配顺序（不使用单字别名子串）：
      1. 精确匹配白名单名称（如“灵山大佛”“灵山胜境”）
      2. 完整别名键精确匹配（如“拈花湾”→“拈花湾禅意小镇”）
      3. 完整白名单名称子串匹配（如“禅意小镇·拈花湾”含“拈花湾”）
    未命中返回 None（该行被丢弃）。
    """
    v = value.strip()
    # 步骤1：精确匹配白名单名称
    if v in names:
        return v
    # 步骤2：完整别名键精确匹配
    if v in aliases:
        return aliases[v]
    # 步骤3：完整白名单名称子串匹配（长名优先）
    for name in sorted(names, key=len, reverse=True):
        if name in v:
            return name
    return None


SECTION_RE = re.compile(r"^[一二三四五六七八九十]+[、．.]")

# 灵山胜境导览手册为整篇叙述文档，标题行为纯文本（无数字编号）：
# 长度短且不含句中标点。据此识别真实小节标题。
_GUIDELINE_HEADER_RE = re.compile(r"[。，！？；]")
# 每条路线下重复出现的通用子标签，需与父级路线标题组合才有意义。
_GENERIC_SUBLABELS = {"讲解重点：", "讲解重点:", "特色体验：", "特色体验:", "餐饮：", "餐饮:", "住宿：", "住宿:"}


def _normalize_route_line(text):
    for prefix in ("路线规划：", "路线规划:", "途经：", "途经:"):
        if text.startswith(prefix) and "→" in text:
            stops = [stop.strip() for stop in text[len(prefix):].split("→") if stop.strip()]
            if len(stops) < 2:
                return text
            start = stops[0]
            if start.endswith("入园"):
                start = start[:-2]
            destination = stops[-1]
            waypoints = stops[1:-1]
            if waypoints:
                return (
                    f"路线：从{start}出发，依次经过{'、'.join(waypoints)}，"
                    f"最后到达{destination}。"
                )
            return f"路线：从{start}出发，最后到达{destination}。"
    return text


def _is_guideline_header(t):
    return len(t) <= 30 and not _GUIDELINE_HEADER_RE.search(t)


def chunk_guideline():
    """guideline.docx 为整篇灵山胜境导览叙述。

    Option A：所有分块 attraction_name 统一为“灵山胜境”（整体景区），
    section 取真实的纯文本小节标题，通用子标签（讲解重点/特色体验等）
    与其父级标题组合，保证 section 有意义、不再误判为“百子戏弥勒”。
    """
    doc = Document(GUIDELINE_DOCX)
    chunks = []
    parent_section = "概述"   # 最近的非通用标题
    current_section = "概述"   # 当前生效小节（可能是 parent - sublabel 组合）
    current_text = []

    def flush():
        if not current_text:
            return
        full = "\n".join(current_text)
        for piece in _split_long_text(full):
            chunks.append({
                "text": piece,
                "metadata": {
                    "source": "guideline",
                    "scenic_area": "灵山胜境",
                    "attraction_name": "灵山胜境",
                    "section": current_section,
                }
            })

    for para in doc.paragraphs:
        t = _normalize_route_line(para.text.strip())
        if not t:
            continue
        if _is_guideline_header(t):
            flush()
            current_text = []
            if t in _GENERIC_SUBLABELS:
                label = t.rstrip("：:")
                current_section = f"{parent_section} - {label}"
            else:
                parent_section = t
                current_section = t
        else:
            current_text.append(t)

    flush()
    return chunks


def chunk_dataset():
    """dataset.docx 为结构化数据集，真实数据存放在两个 Word 表格中：
      表0：灵山胜境（16个子景点 LS-001~016）
      表1：拈花湾禅意小镇（6个子景点 NH-001~006）
    每个子景点组装成一条 chunk，字段拼成可读文本。
    metadata 记录 scenic_area(父景区) / attraction_name(子景点) / attraction_id。
    """
    doc = Document(DATASET_DOCX)
    chunks = []
    for tb in doc.tables:
        rows = tb.rows
        if not rows:
            continue
        header = [c.text.strip() for c in rows[0].cells]
        try:
            i_area = header.index("景区名称")
            i_id   = header.index("景点ID")
            i_name = header.index("景点名称")
        except ValueError:
            continue

        for row in rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < len(header):
                continue
            scenic_area = cells[i_area]
            attraction  = cells[i_name]
            attr_id     = cells[i_id]
            if not attraction:
                continue

            # 拼装：标题行 + 非空字段
            lines = [f"{scenic_area} · {attraction}"]
            for h, v in zip(header, cells):
                if h in ("景区名称", "景点ID", "景点名称") or not v:
                    continue
                lines.append(f"{h}：{v}")
            full = "\n".join(lines)

            for piece in _split_long_text(full):
                chunks.append({
                    "text": piece,
                    "metadata": {
                        "source": "dataset",
                        "scenic_area": scenic_area,
                        "attraction_name": attraction,
                        "attraction_id": attr_id,
                    }
                })
    return chunks


def chunk_xlsx():
    wl = _load_whitelist_full()
    names, aliases = wl["names"], wl["aliases"]
    sub_to_area = wl.get("sub_to_area", {})
    scenic_areas = set(wl.get("scenic_areas", []))
    df = pd.read_excel(XLSX_FILE, dtype=str).fillna("")
    chunks = []

    def _area_of(attraction):
        if attraction in scenic_areas:
            return attraction
        return sub_to_area.get(attraction, attraction)

    # xlsx 为 14 万行游客行为数据，同一景点的 attraction_content 在成百上千行中重复。
    # 先按 (canonical景点名, content) 去重，避免向量库塞入上万条近似重复段落。
    seen = set()
    for _, row in df.iterrows():
        raw_attraction = row.get("attraction_name", "").strip()
        attraction = _match_attraction(raw_attraction, names, aliases)
        if attraction is None:
            continue
        content = row.get("attraction_content", "").strip()
        if not content:
            continue
        # The source spreadsheet contains a mislabeled batch whose
        # attraction_name is 灵山大佛 but whose body describes 嘉兴梅花洲.
        # Requiring the canonical attraction or its parent area in the body
        # prevents a valid-looking label from poisoning the vector index.
        scenic_area = _area_of(attraction)
        if attraction not in content and scenic_area not in content:
            continue
        key = (attraction, content)
        if key in seen:
            continue
        seen.add(key)
        sections = re.split(r"(?=[一二三四五六七八九十]+[、．.])", content)
        for sec in sections:
            sec = sec.strip()
            if not sec:
                continue
            section_title = sec[:20].split("\n")[0]
            for piece in _split_long_text(sec):
                chunks.append({
                    "text": piece,
                    "metadata": {
                        "source": "xlsx",
                        "scenic_area": scenic_area,
                        "attraction_name": attraction,
                        "section": section_title,
                    }
                })
    return chunks


def chunk_extra_documents():
    """Chunk administrator-uploaded DOCX/TXT/Markdown documents.

    Uploaded files are intentionally kept outside the llm repository. Their
    source filename and best-effort attraction metadata remain visible in
    citations, while retrieval still falls back to full-corpus search when no
    attraction can be inferred.
    """
    if not EXTRA_DOCS_DIR.exists():
        return []
    wl = _load_whitelist_full()
    names = wl["names"]
    aliases = wl["aliases"]
    scenic_areas = set(wl.get("scenic_areas", []))
    sub_to_area = wl.get("sub_to_area", {})
    chunks = []
    for path in sorted(EXTRA_DOCS_DIR.iterdir()):
        if not path.is_file() or path.suffix.lower() not in {".docx", ".txt", ".md"}:
            continue
        if path.suffix.lower() == ".docx":
            paragraphs = [p.text.strip() for p in Document(path).paragraphs if p.text.strip()]
        else:
            paragraphs = [
                line.strip()
                for line in path.read_text(encoding="utf-8", errors="ignore").splitlines()
                if line.strip()
            ]
        text = "\n".join(paragraphs)
        if not text:
            continue
        matches = []
        for name in sorted(names, key=len, reverse=True):
            if name in text and not any(name in matched for matched in matches):
                matches.append(name)
        for alias, canonical in aliases.items():
            if alias in text and canonical not in matches:
                matches.append(canonical)
        attraction = matches[0] if len(matches) == 1 else ""
        if attraction in scenic_areas:
            scenic_area = attraction
        else:
            scenic_area = sub_to_area.get(attraction, "")
        for piece in _split_long_text(text):
            chunks.append(
                {
                    "text": piece,
                    "metadata": {
                        "source": f"upload/{path.name}",
                        "scenic_area": scenic_area,
                        "attraction_name": attraction,
                        "section": path.stem,
                    },
                }
            )
    return chunks


def load_all_chunks():
    return chunk_guideline() + chunk_dataset() + chunk_xlsx() + chunk_extra_documents()
