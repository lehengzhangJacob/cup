# scripts/extract_names.py
import json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from rag.config import GUIDELINE_DOCX, DATASET_DOCX, WHITELIST_JSON

# guideline.docx 中提到、但 dataset 表格未单列的景点/别称候选。
EXTRA_CANDIDATES = [
    "灵山梵宫", "五印坛城", "佛手广场", "灵山精舍",
    "三圣殿", "慈恩塔", "灵山佛学院", "拈花塔",
]

# 口语别名 → 规范景点名。
ALIASES = {
    "大佛":   "灵山大佛",
    "梵宫":   "灵山梵宫",
    "坛城":   "五印坛城",
    "精舍":   "灵山精舍",
    "禅寺":   "祥符禅寺",
    "拈花湾": "拈花湾禅意小镇",
}


def extract_paragraph_text(docx_path):
    doc = Document(docx_path)
    return "\n".join(p.text for p in doc.paragraphs)


def extract_from_tables(docx_path):
    """从 dataset.docx 的表格读取权威子景点清单。

    返回 (scenic_areas, sub_names, sub_to_area)。
    """
    doc = Document(docx_path)
    scenic_areas, sub_names, sub_to_area = [], [], {}
    for tb in doc.tables:
        rows = tb.rows
        if not rows:
            continue
        header = [c.text.strip() for c in rows[0].cells]
        try:
            i_area = header.index("景区名称")
            i_name = header.index("景点名称")
        except ValueError:
            continue
        for row in rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) <= max(i_area, i_name):
                continue
            area, name = cells[i_area], cells[i_name]
            if area and area not in scenic_areas:
                scenic_areas.append(area)
            if name and name not in sub_names:
                sub_names.append(name)
                sub_to_area[name] = area
    return scenic_areas, sub_names, sub_to_area


def main():
    scenic_areas, sub_names, sub_to_area = extract_from_tables(DATASET_DOCX)

    # names = 父景区 + 全部子景点 + guideline 中确认存在的额外候选。
    guideline_text = extract_paragraph_text(GUIDELINE_DOCX)
    names = list(scenic_areas) + list(sub_names)
    for extra in EXTRA_CANDIDATES:
        if extra in guideline_text and extra not in names:
            names.append(extra)

    result = {
        "names": names,
        "scenic_areas": scenic_areas,
        "sub_to_area": sub_to_area,
        "aliases": ALIASES,
    }
    WHITELIST_JSON.parent.mkdir(parents=True, exist_ok=True)
    with open(WHITELIST_JSON, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f"scenic_areas={len(scenic_areas)}, sub_attractions={len(sub_names)}, total names={len(names)}", flush=True)
    sys.stdout.buffer.write(("\n".join(names) + "\n").encode("utf-8"))


if __name__ == "__main__":
    main()
